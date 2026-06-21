import os
import uuid
import json
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

def get_storage_dir():
    cwd = Path(os.getcwd())
    storage_path = cwd / 'storage'
    if not storage_path.exists():
        storage_path.mkdir(parents=True, exist_ok=True)
    return storage_path

def document_generator(parameters: dict, player=None):
    action = parameters.get("action", "word")
    content = parameters.get("content", "")
    filename = parameters.get("filename", "")
    data = parameters.get("data", "[]") # Expects JSON string for excel
    
    storage_dir = get_storage_dir()
    
    if not filename:
        filename = f"document_{uuid.uuid4().hex[:8]}"
        
    try:
        if action == "word":
            if not filename.endswith('.docx'):
                filename += '.docx'
            out_path = storage_dir / filename
            
            doc = Document()
            # Basic parsing: treating double newlines as paragraphs
            paragraphs = content.split("\n\n")
            for p in paragraphs:
                if p.startswith("# "):
                    doc.add_heading(p[2:], level=1)
                elif p.startswith("## "):
                    doc.add_heading(p[3:], level=2)
                else:
                    doc.add_paragraph(p)
                    
            doc.save(str(out_path))
            return f"Word document created and saved to {out_path}"
            
        elif action == "excel":
            if not filename.endswith('.xlsx'):
                filename += '.xlsx'
            out_path = storage_dir / filename
            
            wb = Workbook()
            ws = wb.active
            
            try:
                rows = json.loads(data)
                if isinstance(rows, list):
                    for row in rows:
                        if isinstance(row, list):
                            ws.append(row)
                        elif isinstance(row, dict):
                            ws.append(list(row.values()))
            except Exception as e:
                return f"Failed to parse JSON data for Excel: {str(e)}"
                
            wb.save(str(out_path))
            return f"Excel spreadsheet created and saved to {out_path}"
            
        elif action == "pdf":
            if not filename.endswith('.pdf'):
                filename += '.pdf'
            out_path = storage_dir / filename
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Simple text addition
            for line in content.split('\n'):
                # fpdf string encoding workaround for basic ascii
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.cell(200, 10, txt=safe_line, ln=True, align='L')
                
            pdf.output(str(out_path))
            return f"PDF created and saved to {out_path}"
            
        else:
            return f"Unknown action: {action}"
            
    except Exception as e:
        return f"Document Generator Error: {str(e)}"
